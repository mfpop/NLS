"""
Tests for the document upload / import endpoint (api/upload_document.py).

Tests cover:
  - Each extractor function directly (PDF, DOCX, TXT, HTML, escape)
  - The Django view via test client
  - Error handling (missing file, wrong type, oversized file)
"""

import io
import os
import tempfile
from pathlib import Path
from unittest.mock import patch

from django.core.files.uploadedfile import SimpleUploadedFile
from django.test import TestCase, override_settings
from django.test.client import Client

# Module under test
from api.upload_document import (
    _extract_text_pdf,
    _extract_text_docx,
    _extract_text_txt,
    _extract_text_html,
    _escape_html,
)

# ---------------------------------------------------------------------------
#  Helper – build a real in-memory PDF using pdfplumber's write support
# ---------------------------------------------------------------------------

def _make_pdf_bytes(text: str) -> bytes:
    """Create a minimal valid PDF containing *text*."""
    from reportlab.pdfgen import canvas  # shipped with pdfplumber

    buf = io.BytesIO()
    c = canvas.Canvas(buf)
    c.drawString(36, 750, text[:90])  # first 90 chars on page 1
    # If text is long, add a second page
    if len(text) > 90:
        c.showPage()
        c.drawString(36, 750, text[90:180])
    c.save()
    return buf.getvalue()


def _make_docx_bytes() -> bytes:
    """Create a minimal DOCX with a heading, paragraph, bold text, and a table."""
    from docx import Document

    doc = Document()
    doc.add_heading("Test Heading", level=1)
    doc.add_paragraph("This is a test paragraph with some content.")
    p = doc.add_paragraph()
    run = p.add_run("Bold text")
    run.bold = True
    p.add_run(" and normal text.")

    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Header 1"
    table.cell(0, 1).text = "Header 2"
    table.cell(1, 0).text = "Cell A"
    table.cell(1, 1).text = "Cell B"

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ============================================================================
#  Unit tests – individual extractor functions
# ============================================================================

class ExtractTextPdfTests(TestCase):
    """Verify PDF text extraction produces wrapped <p> tags."""

    def test_simple_text(self):
        pdf_bytes = _make_pdf_bytes("Hello PDF world")
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as f:
            f.write(pdf_bytes)
            tmp = f.name
        try:
            html = _extract_text_pdf(tmp)
            self.assertIn("<p>", html)
            self.assertIn("Hello PDF world", html)
            self.assertIn("</p>", html)
        finally:
            os.unlink(tmp)

    def test_html_escaping(self):
        pdf_bytes = _make_pdf_bytes("A < B & C > D")
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as f:
            f.write(pdf_bytes)
            tmp = f.name
        try:
            html = _extract_text_pdf(tmp)
            self.assertIn("&lt;", html)
            self.assertIn("&amp;", html)
            self.assertIn("&gt;", html)
            self.assertNotIn("< B", html)
        finally:
            os.unlink(tmp)

    def test_empty_page_returns_empty(self):
        pdf_bytes = _make_pdf_bytes("")
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as f:
            f.write(pdf_bytes)
            tmp = f.name
        try:
            html = _extract_text_pdf(tmp)
            # Should either be empty or a single empty <p></p>
            self.assertEqual(html, "")
        finally:
            os.unlink(tmp)


class ExtractTextDocxTests(TestCase):
    """Verify DOCX extraction preserves headings, paragraphs, bold, tables, and images."""

    def test_heading_and_paragraph(self):
        docx_bytes = _make_docx_bytes()
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            f.write(docx_bytes)
            tmp = f.name
        try:
            html = _extract_text_docx(tmp)
            self.assertIn("<h1>", html)
            self.assertIn("Test Heading", html)
            self.assertIn("</h1>", html)
            self.assertIn("<p>", html)
            self.assertIn("test paragraph", html)
        finally:
            os.unlink(tmp)

    def test_bold_text(self):
        docx_bytes = _make_docx_bytes()
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            f.write(docx_bytes)
            tmp = f.name
        try:
            html = _extract_text_docx(tmp)
            self.assertIn("<strong>Bold text</strong>", html)
            self.assertIn("normal text", html)
        finally:
            os.unlink(tmp)

    def test_table(self):
        docx_bytes = _make_docx_bytes()
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            f.write(docx_bytes)
            tmp = f.name
        try:
            html = _extract_text_docx(tmp)
            self.assertIn("<table>", html)
            self.assertIn("<th>Header 1</th>", html)
            self.assertIn("<td>Cell A</td>", html)
            self.assertIn("<td>Cell B</td>", html)
        finally:
            os.unlink(tmp)

    def test_inline_image(self):
        """DOCX with embedded inline image produces <img> tag and saved file."""
        docx_bytes = _make_docx_with_image()
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            f.write(docx_bytes)
            tmp = f.name
        try:
            with override_settings(MEDIA_ROOT=tempfile.mkdtemp()):
                html = _extract_text_docx(tmp)
                self.assertIn("<img", html)
                self.assertIn('style="max-width:100%"', html)
                # Verify an image file was saved to disk
                from pathlib import Path
                from django.conf import settings
                img_dir = Path(settings.MEDIA_ROOT) / "editor-images"
                saved_files = list(img_dir.glob("docx_img_*"))
                self.assertGreater(len(saved_files), 0, "No extracted image file found on disk")
        finally:
            os.unlink(tmp)


def _make_png(w: int, h: int, r: int, g: int, b: int) -> bytes:
    """Create a minimal PNG with the given color."""
    import struct
    import zlib

    def _chunk(ctype: bytes, data: bytes) -> bytes:
        c = ctype + data
        return struct.pack(">I", len(data)) + c + struct.pack(">I", zlib.crc32(c) & 0xFFFFFFFF)

    raw_data = b""
    for _ in range(h):
        raw_data += b"\x00" + bytes([r, g, b]) * w
    ihdr = struct.pack(">IIBBBBB", w, h, 8, 2, 0, 0, 0)
    return (
        b"\x89PNG\r\n\x1a\n"
        + _chunk(b"IHDR", ihdr)
        + _chunk(b"IDAT", zlib.compress(raw_data))
        + _chunk(b"IEND", b"")
    )


def _make_docx_with_image() -> bytes:
    """Create a DOCX that contains a 1×1 pixel PNG inline image."""
    from docx import Document
    from docx.shared import Inches

    doc = Document()
    doc.add_paragraph("Before image")

    png_bytes = _make_png(1, 1, 255, 0, 0)
    doc.add_picture(io.BytesIO(png_bytes), width=Inches(1))
    doc.add_paragraph("After image")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


class ExtractTextTxtTests(TestCase):
    """Verify TXT extraction wraps lines in <p> tags."""

    def test_single_paragraph(self):
        html = _extract_text_txt_write("Hello world")
        self.assertEqual(html, "<p>Hello world</p>")

    def test_multiple_paragraphs(self):
        html = _extract_text_txt_write("First para.\n\nSecond para.\n\nThird para.")
        self.assertIn("<p>First para.</p>", html)
        self.assertIn("<p>Second para.</p>", html)
        self.assertIn("<p>Third para.</p>", html)

    def test_html_escaping(self):
        html = _extract_text_txt_write("A < B & C > D")
        self.assertIn("&lt; B &amp; C &gt;", html)

    def test_fallback_line_by_line(self):
        """When there are no double-newlines, the whole text becomes one paragraph with line breaks."""
        html = _extract_text_txt_write("Line one\nLine two\nLine three")
        self.assertEqual(html, "<p>Line one\nLine two\nLine three</p>")


def _extract_text_txt_write(text: str) -> str:
    """Helper: write text to a temp file and call _extract_text_txt."""
    with tempfile.NamedTemporaryFile(suffix=".txt", mode="w", delete=False, encoding="utf-8") as f:
        f.write(text)
        tmp = f.name
    try:
        return _extract_text_txt(tmp)
    finally:
        os.unlink(tmp)


class ExtractTextHtmlTests(TestCase):
    """Verify HTML pass-through."""

    def test_passthrough(self):
        raw = "<h1>Title</h1><p>Some <strong>content</strong>.</p>"
        with tempfile.NamedTemporaryFile(suffix=".html", mode="w", delete=False, encoding="utf-8") as f:
            f.write(raw)
            tmp = f.name
        try:
            html = _extract_text_html(tmp)
            self.assertEqual(html, raw)
        finally:
            os.unlink(tmp)

    def test_bom_stripped(self):
        raw = "\ufeff<p>Hello</p>"
        with tempfile.NamedTemporaryFile(suffix=".html", mode="w", delete=False, encoding="utf-8") as f:
            f.write(raw)
            tmp = f.name
        try:
            html = _extract_text_html(tmp)
            self.assertEqual(html, "<p>Hello</p>")
            self.assertNotIn("\ufeff", html)
        finally:
            os.unlink(tmp)


class EscapeHtmlTests(TestCase):
    """Verify _escape_html handles HTML special characters."""

    def test_ampersand(self):
        self.assertEqual(_escape_html("A & B"), "A &amp; B")

    def test_angle_brackets(self):
        self.assertEqual(_escape_html("<tag>"), "&lt;tag&gt;")

    def test_double_quote(self):
        self.assertEqual(_escape_html('say "hi"'), "say &quot;hi&quot;")

    def test_no_change(self):
        self.assertEqual(_escape_html("plain text"), "plain text")

    def test_combined(self):
        self.assertEqual(
            _escape_html('<a href="x">A & B</a>'),
            "&lt;a href=&quot;x&quot;&gt;A &amp; B&lt;/a&gt;",
        )


# ============================================================================
#  View tests – Django test client
# ============================================================================

class UploadDocumentViewTests(TestCase):
    temp_dir = tempfile.mkdtemp()

    @classmethod
    def tearDownClass(cls):
        import shutil
        shutil.rmtree(cls.temp_dir, ignore_errors=True)
        super().tearDownClass()

    @override_settings(MEDIA_ROOT=temp_dir)
    def setUp(self):
        self.client = Client()
        return super().setUp()
    """Tests for the upload_document Django view via test client."""

    def setUp(self):
        self.client = Client()

    # -- Success cases --

    def test_txt_upload(self):
        file = SimpleUploadedFile("test.txt", b"Hello world\n\nSecond paragraph", content_type="text/plain")
        resp = self.client.post("/api/upload-document/", {"file": file})
        self.assertEqual(resp.status_code, 200)
        data = resp.json()
        self.assertTrue(data["ok"])
        self.assertIn("<p>Hello world</p>", data["html"])
        self.assertIn("<p>Second paragraph</p>", data["html"])
        self.assertEqual(data["file_name"], "test.txt")

    def test_html_upload(self):
        content = b"<h1>Saved</h1><p>Original HTML</p>"
        file = SimpleUploadedFile("doc.html", content, content_type="text/html")
        resp = self.client.post("/api/upload-document/", {"file": file})
        self.assertEqual(resp.status_code, 200)
        data = resp.json()
        self.assertTrue(data["ok"])
        self.assertEqual(data["html"], "<h1>Saved</h1><p>Original HTML</p>")

    def test_docx_upload(self):
        docx_bytes = _make_docx_bytes()
        file = SimpleUploadedFile("report.docx", docx_bytes, content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        resp = self.client.post("/api/upload-document/", {"file": file})
        self.assertEqual(resp.status_code, 200)
        data = resp.json()
        self.assertTrue(data["ok"])
        self.assertIn("<h1>Test Heading</h1>", data["html"])
        self.assertIn("<strong>Bold text</strong>", data["html"])

    def test_pdf_upload(self):
        pdf_bytes = _make_pdf_bytes("PDF content works")
        file = SimpleUploadedFile("doc.pdf", pdf_bytes, content_type="application/pdf")
        resp = self.client.post("/api/upload-document/", {"file": file})
        self.assertEqual(resp.status_code, 200)
        data = resp.json()
        self.assertTrue(data["ok"])
        self.assertIn("PDF content works", data["html"])

    # -- Error cases --

    def test_no_file_returns_400(self):
        resp = self.client.post("/api/upload-document/", {})
        self.assertEqual(resp.status_code, 400)
        data = resp.json()
        self.assertFalse(data["ok"])
        self.assertEqual(data["code"], "FILE_REQUIRED")

    def test_invalid_extension_returns_400(self):
        file = SimpleUploadedFile("data.zip", b"some content", content_type="application/zip")
        resp = self.client.post("/api/upload-document/", {"file": file})
        self.assertEqual(resp.status_code, 400)
        data = resp.json()
        self.assertFalse(data["ok"])
        self.assertEqual(data["code"], "INVALID_TYPE")

    def test_oversized_file_returns_413(self):
        big_content = b"x" * (21 * 1024 * 1024)  # 21 MB
        file = SimpleUploadedFile("big.txt", big_content, content_type="text/plain")
        resp = self.client.post("/api/upload-document/", {"file": file})
        self.assertEqual(resp.status_code, 413)
        data = resp.json()
        self.assertFalse(data["ok"])
        self.assertEqual(data["code"], "FILE_TOO_LARGE")

    @patch("api.upload_document._EXTRACTORS", {".txt": lambda _: (_ for _ in ()).throw(ValueError("Parse explosion"))})
    def test_parse_error_returns_422(self):
        file = SimpleUploadedFile("broken.txt", b"content", content_type="text/plain")
        resp = self.client.post("/api/upload-document/", {"file": file})
        self.assertEqual(resp.status_code, 422)
        data = resp.json()
        self.assertFalse(data["ok"])
        self.assertEqual(data["code"], "PARSE_FAILED")

    def test_wrong_http_method_returns_405(self):
        resp = self.client.get("/api/upload-document/")
        self.assertEqual(resp.status_code, 405)
